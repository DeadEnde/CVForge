"""CVForge API — Vercel serverless function (single self-contained file).

This file contains the parser, themes and portfolio renderer INLINE so the
function has ZERO local imports (no cvforge package, no fastmcp, no path
hacks) — the #1 cause of FUNCTION_INVOCATION_FAILED on Vercel.

It also acts as a catch-all: whatever path Vercel forwards
(/api/health, /health, /api/index, /, ...) is dispatched manually, so routing
ambiguity cannot break the function. GET / serves the landing page (bundled
as landing.html next to this file).

Endpoints:
  GET  /api/health | /health                    -> status
  GET  /api/cv/themes | /cv/themes              -> theme list
  POST /api/cv/parse | /cv/parse                -> CV text -> structured JSON
  POST /api/cv/parse_file | /cv/parse_file      -> upload PDF/DOCX/MD/TXT
  POST /api/cv/generate | /cv/generate          -> CV -> portfolio HTML (EN/AR-RTL)
  GET  / | /index.html                          -> landing page
"""

import json
import re
import uuid
from pathlib import Path

from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import JSONResponse, PlainTextResponse, Response

# ---------------------------------------------------------------------------
# MINI CV PARSER (ported from cvforge/cv_parser.py, self-contained)
# ---------------------------------------------------------------------------
DOMAIN_KEYWORDS = {
    "developer": ["developer",
                 "software",
                 "engineer",
                 "programmer",
                 "fullstack",
                 "full stack",
                 "backend",
                 "frontend",
                 "front-end",
                 "react",
                 "vue",
                 "angular",
                 "python",
                 "javascript",
                 "typescript",
                 "node",
                 "java",
                 "php",
                 "laravel",
                 "web developer",
                 "api",
                 "sql",
                 "android",
                 "ios"],
    "designer": ["designer",
                 "design",
                 "ui",
                 "ux",
                 "ui/ux",
                 "figma",
                 "photoshop",
                 "illustrator",
                 "branding",
                 "brand identity",
                 "logo",
                 "graphic",
                 "visual",
                 "creative",
                 "art director",
                 "typography",
                 "prototype",
                 "wireframe",
                 "interaction design"],
    "photographer": ["photograph",
                 "photographer",
                 "photo",
                 "camera",
                 "lighting",
                 "editing",
                 "lightroom",
                 "retouch",
                 "studio",
                 "wedding"],
    "marketer": ["marketing",
                 "seo",
                 "sem",
                 "social media",
                 "growth",
                 "ads",
                 "campaign",
                 "brand manager",
                 "email marketing",
                 "crm",
                 "content"],
    "data": ["data analyst",
                 "bi ",
                 "analytics",
                 "power bi",
                 "tableau",
                 "statistic",
                 "pandas",
                 "dashboard",
                 "etl",
                 "sql"],
    "manager": ["manager",
                 "director",
                 "project manager",
                 "product manager",
                 "product owner",
                 "chef de projet",
                 "scrum",
                 "agile",
                 "leadership",
                 "team lead"],
    "teacher": ["teacher",
                 "professor",
                 "professeur",
                 "trainer",
                 "formation",
                 "education",
                 "tutor",
                 "lecturer",
                 "pedagog"],
    "finance": ["finance",
                 "financial",
                 "banking",
                 "banque",
                 "fiscal",
                 "comptabilité",
                 "investor",
                 "comptable"],
    "health": ["doctor",
                 "nurse",
                 "infirmier",
                 "medecin",
                 "médecin",
                 "pharmac",
                 "dentist",
                 "physiotherap",
                 "therapist",
                 "clinician",
                 "caregiver"],
    "engineer": ["civil engineer",
                 "mechanical engineer",
                 "electrical engineer",
                 "ingénieur",
                 "ingenieur",
                 "construction",
                 "structural",
                 "mep",
                 "site engineer",
                 "project engineer"],
    "generic": [],
    "ai": ["artificial intelligence",
                 "ai engineer",
                 "machine learning",
                 "deep learning",
                 "nlp",
                 "computer vision",
                 "llm",
                 "ml engineer",
                 "data scientist",
                 "chatbot",
                 "neural",
                 "generative ai",
                 "prompt engineer",
                 "ai model"],
    "cybersecurity": ["cybersecurity",
                 "cyber security",
                 "security analyst",
                 "penetration",
                 "pen-test",
                 "ethical hacker",
                 "infosec",
                 "it security",
                 "network security",
                 "soc analyst",
                 "cissp",
                 "vulnerability"],
    "cloud": ["cloud",
                 "aws",
                 "azure",
                 "gcp",
                 "cloud engineer",
                 "cloud architect",
                 "serverless"],
    "game": ["game developer",
                 "gamedev",
                 "unity",
                 "unreal",
                 "game design",
                 "game engine",
                 "esports",
                 "gaming",
                 "level design"],
    "blockchain": ["blockchain",
                 "web3",
                 "solidity",
                 "smart contract",
                 "crypto",
                 "bitcoin",
                 "ethereum",
                 "defi",
                 "nft",
                 "token"],
    "mobile": ["mobile developer",
                 "ios developer",
                 "android developer",
                 "react native",
                 "flutter developer",
                 "swift",
                 "kotlin",
                 "mobile app",
                 "app developer"],
    "devops": ["devops",
                 "site reliability",
                 "sre",
                 "ci/cd",
                 "ci cd",
                 "docker",
                 "kubernetes",
                 "terraform",
                 "ansible",
                 "jenkins",
                 "infrastructure",
                 "platform engineer",
                 "cloudops"],
    "robotics": ["robotics",
                 "robot",
                 "automation engineer",
                 "embedded",
                 "iot",
                 "arduino",
                 "raspberry",
                 "mechatronics",
                 "automation"],
    "support": ["customer support",
                 "support specialist",
                 "helpdesk",
                 "service desk",
                 "customer service",
                 "call center",
                 "technical support",
                 "customer care",
                 "ticket"],
    "economist": ["economist",
                 "economics",
                 "econometrics",
                 "macroeconomic",
                 "microeconomic",
                 "economic analyst",
                 "policy analyst",
                 "development economics"],
    "accounting": ["accountant",
                 "accounting",
                 "audit",
                 "auditor",
                 "tax advisor",
                 "bookkeeper",
                 "cpa",
                 "controller",
                 "payroll",
                 "comptable",
                 "fiscaliste",
                 "comptabilité"],
    "ecommerce": ["ecommerce",
                 "e-commerce",
                 "shopify",
                 "dropshipping",
                 "dropship",
                 "amazon seller",
                 "marketplace",
                 "online store",
                 "store manager",
                 "retail",
                 "merchandising",
                 "woocommerce"],
    "business": ["business",
                 "operations",
                 "operations manager",
                 "corporate",
                 "strategy",
                 "b2b",
                 "b2c"],
    "founder": ["founder",
                 "co-founder",
                 "cofounder",
                 "entrepreneur",
                 "startup",
                 "ceo",
                 "owner",
                 "self-employed",
                 "solo founder"],
    "sales": ["sales",
                 "account executive",
                 "business development",
                 "sdr",
                 "bdr",
                 "sales manager",
                 "sales representative",
                 "commercial",
                 "salesforce",
                 "closer"],
    "consulting": ["consultant",
                 "consulting",
                 "management consulting",
                 "strategy consultant",
                 "advisory"],
    "logistics": ["logistics",
                 "supply chain",
                 "warehouse",
                 "procurement",
                 "freight",
                 "transportation",
                 "customs",
                 "inventory",
                 "distribution"],
    "realestate": ["real estate",
                 "realtor",
                 "property agent",
                 "real estate agent",
                 "leasing",
                 "landlord",
                 "property management"],
    "hr": ["human resources",
                 "hr manager",
                 "recruiter",
                 "recruitment",
                 "talent acquisition",
                 "people operations",
                 "headhunter",
                 "hrbp",
                 "staffing"],
    "doctor": ["doctor",
                 "physician",
                 "medical doctor",
                 "surgeon",
                 "general practitioner",
                 "dentist",
                 "pharmacist",
                 "nurse",
                 "clinician",
                 "pediatric",
                 "gynecolog",
                 "cardiolog",
                 "radiolog",
                 "anesthesi",
                 "medecin",
                 "médecin",
                 "paramedic",
                 "physiotherap",
                 "infirmier"],
    "psychology": ["psychologist",
                 "psychology",
                 "psychiatry",
                 "psychiatrist",
                 "counselor",
                 "counsellor",
                 "mental health",
                 "psychotherapy",
                 "psychologue"],
    "nutrition": ["nutritionist",
                 "nutrition",
                 "dietitian",
                 "dietetics",
                 "diet coach",
                 "nutritionniste"],
    "fitness": ["personal trainer",
                 "fitness",
                 "gym",
                 "crossfit",
                 "athlete",
                 "sports coach",
                 "strength coach",
                 "yoga instructor",
                 "pilates"],
    "veterinary": ["veterinarian",
                 "veterinary",
                 "animal hospital",
                 "vet clinic",
                 "vétérinaire"],
    "uiux": ["ui/ux",
                 "ux designer",
                 "ui designer",
                 "product designer",
                 "interaction designer",
                 "usability",
                 "user research",
                 "ux research",
                 "ux engineer"],
    "interior": ["interior designer",
                 "interior design",
                 "space planning",
                 "interior architecture",
                 "decorator",
                 "décorateur"],
    "fashion": ["fashion",
                 "stylist",
                 "fashion designer",
                 "apparel",
                 "clothing",
                 "couture",
                 "modest fashion"],
    "videographer": ["videographer",
                 "video editor",
                 "filmmaker",
                 "cinematographer",
                 "video production",
                 "video creator"],
    "musician": ["musician",
                 "music producer",
                 "singer",
                 "songwriter",
                 "composer",
                 "dj",
                 "sound engineer",
                 "guitarist",
                 "pianist",
                 "band"],
    "artist": ["artist",
                 "fine art",
                 "fine artist",
                 "sculptor",
                 "tattoo artist",
                 "ceramics",
                 "illustration art"],
    "writer": ["writer",
                 "copywriter",
                 "author",
                 "novelist",
                 "content writer",
                 "technical writer",
                 "scriptwriter",
                 "essayist",
                 "poet",
                 "rédacteur"],
    "journalist": ["journalist",
                 "journalism",
                 "reporter",
                 "news editor",
                 "press",
                 "editor in chief",
                 "correspondent",
                 "media"],
    "translator": ["translator",
                 "translation",
                 "interpreter",
                 "localization",
                 "linguist",
                 "traducteur",
                 "linguistics"],
    "contentcreator": ["content creator",
                 "influencer",
                 "youtuber",
                 "streamer",
                 "blogger",
                 "podcast host",
                 "instagram",
                 "tiktok",
                 "creator economy",
                 "community manager"],
    "architecture": ["architect",
                 "architecture",
                 "architectural",
                 "urban planning",
                 "urban planner",
                 "landscape architect"],
    "researcher": ["researcher",
                 "research scientist",
                 "postdoc",
                 "phd candidate",
                 "phd student",
                 "academic researcher",
                 "laboratory research"],
    "scientist": ["scientist",
                 "physics",
                 "chemistry",
                 "biology",
                 "biotech",
                 "biotechnology",
                 "geology",
                 "astronomy",
                 "genetics",
                 "laboratory",
                 "microbiology",
                 "neuroscience"],
    "historian": ["historian",
                 "history",
                 "archivist",
                 "museum",
                 "heritage",
                 "archaeolog"],
    "chef": ["chef",
                 "cook",
                 "culinary",
                 "pastry chef",
                 "sous chef",
                 "head chef",
                 "kitchen",
                 "gastronomy",
                 "baker",
                 "cuisinier"],
    "hospitality": ["hospitality",
                 "hotel",
                 "waiter",
                 "waitress",
                 "bartender",
                 "barista",
                 "hostess",
                 "concierge",
                 "restaurant manager",
                 "front office",
                 "hôtellerie"],
    "tourism": ["tourism",
                 "tour guide",
                 "travel consultant",
                 "travel agency",
                 "destination",
                 "travel advisor"],
    "trades": ["electrician",
                 "plumber",
                 "carpenter",
                 "mechanic",
                 "welder",
                 "hvac",
                 "technician",
                 "handyman",
                 "mason",
                 "locksmith",
                 "glazier",
                 "tiler",
                 "auto mechanic",
                 "technicien"],
    "beauty": ["hairdresser",
                 "hair stylist",
                 "barber",
                 "beautician",
                 "esthetician",
                 "makeup artist",
                 "nail technician",
                 "salon",
                 "spa",
                 "cosmetology",
                 "coiffeur",
                 "manicure"],
    "aviation": ["pilot",
                 "aviation",
                 "flight attendant",
                 "aircraft",
                 "air traffic",
                 "captain",
                 "stewardess",
                 "airline",
                 "drone pilot"],
    "agriculture": ["farmer",
                 "agriculture",
                 "agronomy",
                 "agronomist",
                 "vineyard",
                 "harvest",
                 "livestock",
                 "greenhouse",
                 "horticultur",
                 "farming"],
    "social": ["social worker",
                 "nonprofit",
                 "non-profit",
                 "ngo",
                 "charity",
                 "humanitarian",
                 "community outreach",
                 "case manager",
                 "community development",
                 "volunteer"],
    "government": ["government",
                 "public administration",
                 "civil servant",
                 "politician",
                 "diplomacy",
                 "public policy",
                 "ministry",
                 "municipal",
                 "diplomat"],
    "legal": ["lawyer",
                 "attorney",
                 "legal",
                 "counsel",
                 "judge",
                 "notary",
                 "paralegal",
                 "barrister",
                 "solicitor",
                 "jurist",
                 "compliance",
                 "magistrat",
                 "avocat",
                 "droit"],
}

DOMAIN_LABELS = {
    "developer": "Software Developer",
    "designer": "Creative Designer",
    "photographer": "Photographer",
    "marketer": "Marketing & Growth",
    "data": "Data Professional",
    "manager": "Product / Project Manager",
    "teacher": "Educator",
    "finance": "Finance Professional",
    "health": "Healthcare Professional",
    "engineer": "Engineer",
    "generic": "Professional",
    "ai": "AI / Machine Learning Engineer",
    "cybersecurity": "Cybersecurity Professional",
    "cloud": "Cloud Engineer",
    "game": "Game Developer",
    "blockchain": "Blockchain Developer",
    "mobile": "Mobile Developer",
    "devops": "DevOps / SRE Engineer",
    "robotics": "Robotics & Automation Engineer",
    "support": "Customer Support Specialist",
    "economist": "Economist",
    "accounting": "Accountant / Auditor",
    "ecommerce": "E-commerce Specialist",
    "business": "Business & Operations",
    "founder": "Founder / Entrepreneur",
    "sales": "Sales Professional",
    "consulting": "Business Consultant",
    "logistics": "Logistics & Supply Chain",
    "realestate": "Real Estate Professional",
    "hr": "HR / Talent Acquisition",
    "doctor": "Medical Professional",
    "psychology": "Psychologist / Therapist",
    "nutrition": "Nutritionist",
    "fitness": "Fitness & Sports Coach",
    "veterinary": "Veterinarian",
    "uiux": "Product / UX Designer",
    "interior": "Interior Designer",
    "fashion": "Fashion & Styling",
    "videographer": "Video & Film Maker",
    "musician": "Musician / Producer",
    "artist": "Visual Artist",
    "writer": "Writer / Copywriter",
    "journalist": "Journalist",
    "translator": "Translator / Interpreter",
    "contentcreator": "Content Creator",
    "architecture": "Architect",
    "researcher": "Researcher",
    "scientist": "Scientist",
    "historian": "Historian",
    "chef": "Chef / Culinary Artist",
    "hospitality": "Hospitality Professional",
    "tourism": "Tourism & Travel",
    "trades": "Skilled Trades Professional",
    "beauty": "Beauty & Wellness",
    "aviation": "Aviation Professional",
    "agriculture": "Agriculture Professional",
    "social": "Social Impact Professional",
    "government": "Public Sector Professional",
    "legal": "Legal Professional",
}

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+?\d[\d\s.\-()]{7,})")
URL_RE = re.compile(r"(?:https?://|www\.)[^\s\"'<>]+|github\.com/[\w-]+|linkedin\.com/in/[\w-]+")
SECTION_RE = {
    "experience": r"(experience|expérience|work history|parcours professionnel|employment)",
    "education": r"(education|formation|études|academic|diploma|degree)",
    "skills": r"(skills|compétences|technical skills|technologies|expertise)",
    "projects": r"(projects|projets|portfolio|réalisations|selected works)",
    "summary": r"(summary|profile|about|à propos|profil|objective|objectif)",
    "languages": r"(languages|langues)",
}
DATE_RANGE_RE = re.compile(
    r"\(\s*(?:19|20)\d{2}\s*[–—-]\s*(?:(?:19|20)\d{2}|present|présent|now|today|en\s+cours|aujourd)",
    re.IGNORECASE,
)
HEADER_RE = re.compile(r"^[A-ZÀ-Ý0-9][\w\s'’\-–—&/+,.()%£€:]{2,160}$")


def detect_domain(text: str) -> str:
    t = text.lower()
    best, best_score = "generic", 0
    for domain, kws in DOMAIN_KEYWORDS.items():
        score = sum(t.count(kw) for kw in kws)
        if score > best_score:
            best, best_score = domain, score
    return best


def _split_sections(text: str) -> dict:
    sections, current = {}, None
    for raw in text.splitlines():
        s = re.sub(r"^#{1,6}\s*", "", raw.strip()).strip()
        if not s:
            continue
        matched = None
        for key, pattern in SECTION_RE.items():
            if re.match(rf"^{pattern}\s*:?\s*$", s, re.IGNORECASE):
                matched = key
                break
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        if current:
            sections[current].append(s)
    for k in sections:
        sections[k] = "\n".join(sections[k])
    return sections


def _simple_parse(lines: list[str], joined: str) -> dict:
    name, title = "", ""
    for line in lines[:6]:
        if (re.fullmatch(r"[A-ZÀ-Ýa-zà-ÿ'’\-–]+(?:\s+[A-ZÀ-Ýa-zà-ÿ'’\-–]+){0,3}", line)
                and len(line) < 60 and not EMAIL_RE.search(line)):
            name = line
            break
    for line in lines[1:9]:
        if (line != name and len(line) < 80 and not EMAIL_RE.search(line)
                and not PHONE_RE.search(line)
                and not re.match(r"^(summary|profile|experience|skills|education|projects|about)", line, re.I)):
            title = line
            break
    email = EMAIL_RE.search(joined)
    phone = PHONE_RE.search(joined)
    urls = URL_RE.findall(joined)
    location = ""
    loc_re = re.compile(r"^[A-ZÀ-Ý][\w'’\- ]+,\s*[A-ZÀ-Ý][\w'’\- ]+$")
    for line in lines[:10]:
        for seg in [x.strip() for x in line.split("|")]:
            if loc_re.match(seg) and not EMAIL_RE.search(seg) and "http" not in seg:
                location = seg
                break
        if location:
            break
    return {"name": name, "title": title,
            "email": email.group(0) if email else "",
            "phone": phone.group(0).strip() if phone else "",
            "website": urls[0] if urls else "", "location": location}


def _bullets(block: str, each_bullet_is_entry: bool = False) -> list[dict]:
    items, current = [], None
    for raw in (block or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        is_bullet = bool(re.match(r"^[-•*·◦▪]\s", line))
        content = re.sub(r"^[-•*·◦▪]\s*", "", line).strip()
        if not content:
            continue
        if each_bullet_is_entry and is_bullet:
            items.append({"title": content, "details": []})
            continue
        is_header = bool(DATE_RANGE_RE.search(content)) or (
            not is_bullet and len(content) < 120 and HEADER_RE.match(content)
            and not content.endswith("."))
        if is_header:
            if current is not None and (current["title"] or current["details"]):
                items.append(current)
            current = {"title": content, "details": []}
            continue
        if is_bullet:
            if current is None:
                current = {"title": content, "details": []}
            else:
                current["details"].append(content)
            continue
        if current is None:
            current = {"title": content[:80], "details": []}
        else:
            current["details"].append(content)
    if current is not None and (current["title"] or current["details"]):
        items.append(current)
    return items


def parse_cv(text: str) -> dict:
    if len(text) < 400 and Path(text).exists():
        path = Path(text)
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                from pypdf import PdfReader
                text = "\n".join(p.extract_text() or "" for p in PdfReader(str(path)).pages)
            elif suffix == ".docx":
                import docx
                text = "\n".join(p.text for p in docx.Document(str(path)).paragraphs)
            else:
                text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return parse_cv(path.read_text(encoding="utf-8", errors="ignore"))

    lines = [re.sub(r"^#{1,6}\s*", "", l.strip()).strip() for l in text.splitlines() if l.strip()]
    joined = "\n".join(lines)
    base = _simple_parse(lines, joined)
    sections = _split_sections(text)
    domain = detect_domain(text)

    skills = []
    for line in (sections.get("skills") or "").splitlines():
        line = line.lstrip("-•*·,;").strip()
        if not line:
            continue
        if line.count(",") >= 1 and not line.startswith(("C++", "C#")):
            skills.extend([p for p in (s.strip() for s in line.split(",")) if p][:8])
        else:
            skills.append(line)
    skills = skills[:24]

    languages = [l.strip().lstrip("-•*·,").strip()
                 for l in (sections.get("languages") or "").splitlines() if l.strip()]
    if len(languages) == 1 and "," in languages[0]:
        languages = [x.strip() for x in languages[0].split(",") if x.strip()]

    return {
        "name": base["name"], "title": base["title"], "email": base["email"],
        "phone": base["phone"], "location": base["location"], "website": base["website"],
        "summary": (sections.get("summary") or "")[:600],
        "skills": skills,
        "experience": _bullets(sections.get("experience") or "", False)[:8],
        "education": _bullets(sections.get("education") or "", False)[:6],
        "projects": _bullets(sections.get("projects") or "", True)[:8],
        "languages": languages[:8],
        "domain": domain, "domain_label": DOMAIN_LABELS.get(domain, "Professional"),
    }


# ---------------------------------------------------------------------------
# THEMES (ported from cvforge/themes.py)
# ---------------------------------------------------------------------------
THEMES = {
    "developer": {"name": "Neon Code", "pal": {"bg": "#070b14", "bg2": "#0b1220", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#6366f1", "accent2": "#22d3ee", "text": "#e7ecf6", "muted": "#8b93a7", "bt": "#fff"}, "grad": ["#6366f1", "#22d3ee", "#a855f7"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "designer": {"name": "Studio Rose", "pal": {"bg": "#140a12", "bg2": "#1b0e18", "card": "rgba(255,255,255,0.06)", "line": "rgba(255,255,255,0.09)", "accent": "#f472b6", "accent2": "#fb923c", "text": "#f8edf4", "muted": "#a3909e", "bt": "#fff"}, "grad": ["#f472b6", "#fb923c", "#facc15"], "font": "Georgia,'Times New Roman',serif", "radius": "28px"},
    "photographer": {"name": "Golden Hour", "pal": {"bg": "#0d0b09", "bg2": "#14110c", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#f59e0b", "accent2": "#fde68a", "text": "#f5efe4", "muted": "#9c9284", "bt": "#fff"}, "grad": ["#f59e0b", "#fbbf24", "#78716c"], "font": "Georgia,'Times New Roman',serif", "radius": "12px"},
    "marketer": {"name": "Growth Green", "pal": {"bg": "#06120e", "bg2": "#0a1a14", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#34d399", "accent2": "#4ade80", "text": "#eaf6ef", "muted": "#8aa79a", "bt": "#fff"}, "grad": ["#34d399", "#4ade80", "#a3e635"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "20px"},
    "data": {"name": "Data Pulse", "pal": {"bg": "#080a12", "bg2": "#0c101e", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#818cf8", "accent2": "#2dd4bf", "text": "#e8ecf7", "muted": "#8a92a8", "bt": "#fff"}, "grad": ["#6366f1", "#2dd4bf", "#0ea5e9"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "16px"},
    "manager": {"name": "Executive", "pal": {"bg": "#0c0d10", "bg2": "#12141a", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#f8fafc", "accent2": "#94a3b8", "text": "#f1f5f9", "muted": "#94a3b8", "bt": "#fff"}, "grad": ["#e2e8f0", "#94a3b8", "#475569"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "teacher": {"name": "Scholar", "pal": {"bg": "#0d1014", "bg2": "#131820", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#60a5fa", "accent2": "#fbbf24", "text": "#eef2f7", "muted": "#8f9aad", "bt": "#fff"}, "grad": ["#60a5fa", "#fbbf24", "#34d399"], "font": "Georgia,'Times New Roman',serif", "radius": "14px"},
    "finance": {"name": "Fintech", "pal": {"bg": "#0a0f0c", "bg2": "#101712", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#10b981", "accent2": "#facc15", "text": "#ecf5ef", "muted": "#8fa295", "bt": "#fff"}, "grad": ["#10b981", "#facc15", "#34d399"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "12px"},
    "health": {"name": "Care", "pal": {"bg": "#0a0f12", "bg2": "#0f161b", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#38bdf8", "accent2": "#f472b6", "text": "#eaf3f8", "muted": "#8aa0ad", "bt": "#fff"}, "grad": ["#38bdf8", "#f472b6", "#a5b4fc"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "engineer": {"name": "Builder", "pal": {"bg": "#0b0e10", "bg2": "#121619", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#f97316", "accent2": "#fbbf24", "text": "#f2f4f6", "muted": "#96a0a8", "bt": "#fff"}, "grad": ["#f97316", "#fbbf24", "#ef4444"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "generic": {"name": "Aurora", "pal": {"bg": "#0a0c14", "bg2": "#10131f", "card": "rgba(255,255,255,0.05)", "line": "rgba(255,255,255,0.09)", "accent": "#a78bfa", "accent2": "#f472b6", "text": "#eef0f8", "muted": "#8d93a8", "bt": "#fff"}, "grad": ["#a78bfa", "#f472b6", "#38bdf8"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "20px"},
    "minimal": {"name": "Snow Minimal", "pal": {"bg": "#f5f7fb", "bg2": "#ffffff", "card": "rgba(255,255,255,0.8)", "line": "rgba(15,23,42,0.08)", "accent": "#0ea5e9", "accent2": "#6366f1", "text": "#0f172a", "muted": "#64748b", "bt": "#ffffff"}, "grad": ["#0ea5e9", "#6366f1", "#8b5cf6"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "14px"},
    "classic": {"name": "Paper & Ink", "pal": {"bg": "#f6f1e7", "bg2": "#fdfaf3", "card": "rgba(255,253,247,0.8)", "line": "rgba(41,37,36,0.12)", "accent": "#b45309", "accent2": "#57534e", "text": "#292524", "muted": "#78716c", "bt": "#ffffff"}, "grad": ["#c2410c", "#b45309", "#78716c"], "font": "Georgia,'Times New Roman',serif", "radius": "6px"},
    "corporate": {"name": "Corporate Pro", "pal": {"bg": "#f8fafc", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(15,23,42,0.08)", "accent": "#1d4ed8", "accent2": "#0ea5e9", "text": "#0f172a", "muted": "#64748b", "bt": "#ffffff"}, "grad": ["#1d4ed8", "#3b82f6", "#0ea5e9"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "pastel": {"name": "Playful Pastel", "pal": {"bg": "#fdf2f8", "bg2": "#fefce8", "card": "rgba(255,255,255,0.75)", "line": "rgba(76,29,149,0.10)", "accent": "#ec4899", "accent2": "#8b5cf6", "text": "#3d3a5c", "muted": "#8d84a8", "bt": "#ffffff"}, "grad": ["#f472b6", "#a78bfa", "#60a5fa"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "24px"},
    "cyber": {"name": "Neon Cyber", "pal": {"bg": "#050014", "bg2": "#0a0220", "card": "rgba(255,79,216,0.06)", "line": "rgba(255,79,216,0.22)", "accent": "#ff2fd6", "accent2": "#00f0ff", "text": "#f2eaff", "muted": "#9d8fd0", "bt": "#fff"}, "grad": ["#ff2fd6", "#00f0ff", "#7c3aed"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "6px"},
    "retro": {"name": "Retro Terminal", "pal": {"bg": "#0a0f0a", "bg2": "#0d140d", "card": "rgba(0,255,65,0.05)", "line": "rgba(0,255,65,0.20)", "accent": "#00ff41", "accent2": "#39ff14", "text": "#d9ffdd", "muted": "#6da372", "bt": "#fff"}, "grad": ["#00ff41", "#39ff14", "#4ade80"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "4px"},
    "gold": {"name": "Luxury Gold", "pal": {"bg": "#0c0a07", "bg2": "#14110a", "card": "rgba(255,255,255,0.04)", "line": "rgba(234,179,8,0.25)", "accent": "#eab308", "accent2": "#fef3c7", "text": "#faf6ea", "muted": "#a89b84", "bt": "#1a1206"}, "grad": ["#b8860b", "#eab308", "#fef3c7"], "font": "Georgia,'Times New Roman',serif", "radius": "4px"},
    "glass": {"name": "Liquid Glass", "pal": {"bg": "#070b1e", "bg2": "#0b1030", "card": "rgba(255,255,255,0.09)", "line": "rgba(255,255,255,0.16)", "accent": "#5eead4", "accent2": "#a5b4fc", "text": "#f0f4ff", "muted": "#9aa7d8", "bt": "#fff"}, "grad": ["#22d3ee", "#818cf8", "#f472b6"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "26px"},
    "noir": {"name": "Noir Mono", "pal": {"bg": "#0a0a0a", "bg2": "#121212", "card": "rgba(255,255,255,0.06)", "line": "rgba(255,255,255,0.14)", "accent": "#ffffff", "accent2": "#a3a3a3", "text": "#f5f5f5", "muted": "#9ca3af", "bt": "#fff"}, "grad": ["#525252", "#737373", "#a3a3a3"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "2px"},
    "ocean": {"name": "Deep Ocean", "pal": {"bg": "#041220", "bg2": "#07223a", "card": "rgba(255,255,255,0.05)", "line": "rgba(34,211,238,0.22)", "accent": "#22d3ee", "accent2": "#60a5fa", "text": "#e8f6ff", "muted": "#8fb0cc", "bt": "#fff"}, "grad": ["#0ea5e9", "#22d3ee", "#3b82f6"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "22px"},
    "forest": {"name": "Deep Forest", "pal": {"bg": "#07130c", "bg2": "#0c1e13", "card": "rgba(255,255,255,0.05)", "line": "rgba(74,222,128,0.20)", "accent": "#34d399", "accent2": "#a3e635", "text": "#eaf7ee", "muted": "#93b5a0", "bt": "#fff"}, "grad": ["#10b981", "#84cc16", "#22c55e"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "royal": {"name": "Royal Violet", "pal": {"bg": "#0d0718", "bg2": "#170b2e", "card": "rgba(255,255,255,0.05)", "line": "rgba(167,139,250,0.25)", "accent": "#a78bfa", "accent2": "#f0abfc", "text": "#f4efff", "muted": "#ab9bd8", "bt": "#fff"}, "grad": ["#7c3aed", "#a78bfa", "#f0abfc"], "font": "Georgia,'Times New Roman',serif", "radius": "22px"},
    "ai": {"name": "Neural AI", "pal": {"bg": "#060913", "bg2": "#0a0f22", "card": "rgba(255,255,255,0.05)", "line": "rgba(139,92,246,0.28)", "accent": "#8b5cf6", "accent2": "#22d3ee", "text": "#ece9ff", "muted": "#9a94c9", "bt": "#fff"}, "grad": ["#8b5cf6", "#22d3ee", "#6366f1"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "16px"},
    "cybersecurity": {"name": "Zero Day", "pal": {"bg": "#070a10", "bg2": "#0c1018", "card": "rgba(255,255,255,0.05)", "line": "rgba(34,197,94,0.25)", "accent": "#22c55e", "accent2": "#ef4444", "text": "#e8f5ec", "muted": "#8fa596", "bt": "#052e16"}, "grad": ["#22c55e", "#16a34a", "#ef4444"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "8px"},
    "cloud": {"name": "Cloud Nine", "pal": {"bg": "#081120", "bg2": "#0c1a30", "card": "rgba(255,255,255,0.05)", "line": "rgba(56,189,248,0.25)", "accent": "#38bdf8", "accent2": "#a5b4fc", "text": "#eaf4ff", "muted": "#8fa3c2", "bt": "#fff"}, "grad": ["#38bdf8", "#60a5fa", "#22d3ee"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "22px"},
    "game": {"name": "Pixel Arcade", "pal": {"bg": "#150a2b", "bg2": "#1b1038", "card": "rgba(255,255,255,0.06)", "line": "rgba(244,114,182,0.3)", "accent": "#f472b6", "accent2": "#fbbf24", "text": "#f7ecff", "muted": "#ab93c9", "bt": "#fff"}, "grad": ["#f472b6", "#fbbf24", "#8b5cf6"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "blockchain": {"name": "Chain Block", "pal": {"bg": "#0d0a06", "bg2": "#14100a", "card": "rgba(255,255,255,0.05)", "line": "rgba(245,158,11,0.28)", "accent": "#f59e0b", "accent2": "#fde68a", "text": "#fbf3e4", "muted": "#a3967d", "bt": "#1a1206"}, "grad": ["#b45309", "#f59e0b", "#fde68a"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "8px"},
    "mobile": {"name": "App Craft", "pal": {"bg": "#f3f6ff", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(37,99,235,0.14)", "accent": "#2563eb", "accent2": "#22d3ee", "text": "#0f1b3d", "muted": "#64748b", "bt": "#fff"}, "grad": ["#2563eb", "#38bdf8", "#6366f1"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "20px"},
    "devops": {"name": "Ops Line", "pal": {"bg": "#0b0f14", "bg2": "#10161d", "card": "rgba(255,255,255,0.05)", "line": "rgba(249,115,22,0.25)", "accent": "#f97316", "accent2": "#38bdf8", "text": "#eef4f8", "muted": "#94a3b0", "bt": "#fff"}, "grad": ["#f97316", "#38bdf8", "#a3e635"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "6px"},
    "robotics": {"name": "Robotics Lab", "pal": {"bg": "#07100e", "bg2": "#0c1815", "card": "rgba(255,255,255,0.05)", "line": "rgba(45,212,191,0.25)", "accent": "#2dd4bf", "accent2": "#fb7185", "text": "#e9f7f3", "muted": "#8aa79f", "bt": "#fff"}, "grad": ["#2dd4bf", "#14b8a6", "#fb7185"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "support": {"name": "Helpdesk", "pal": {"bg": "#f0f7ff", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(2,132,199,0.16)", "accent": "#0284c7", "accent2": "#14b8a6", "text": "#0c1c2e", "muted": "#5f7488", "bt": "#fff"}, "grad": ["#0ea5e9", "#14b8a6", "#6366f1"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "16px"},
    "economist": {"name": "Economist", "pal": {"bg": "#f7f3e7", "bg2": "#fffdf6", "card": "rgba(255,255,255,0.8)", "line": "rgba(26,93,58,0.14)", "accent": "#1a5d3a", "accent2": "#b45309", "text": "#1c2a20", "muted": "#6c7a6a", "bt": "#fff"}, "grad": ["#1a5d3a", "#b45309", "#78716c"], "font": "Georgia,'Times New Roman',serif", "radius": "8px"},
    "accounting": {"name": "Ledger", "pal": {"bg": "#f6faf9", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(5,150,105,0.16)", "accent": "#059669", "accent2": "#0f766e", "text": "#0f2b22", "muted": "#5f7d70", "bt": "#fff"}, "grad": ["#059669", "#0f766e", "#84cc16"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "8px"},
    "ecommerce": {"name": "Storefront", "pal": {"bg": "#fff6ec", "bg2": "#fffdf9", "card": "rgba(255,255,255,0.85)", "line": "rgba(234,88,12,0.16)", "accent": "#ea580c", "accent2": "#db2777", "text": "#331b10", "muted": "#8a7868", "bt": "#fff"}, "grad": ["#ea580c", "#f59e0b", "#db2777"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "business": {"name": "Boardroom", "pal": {"bg": "#0b1120", "bg2": "#101a33", "card": "rgba(255,255,255,0.05)", "line": "rgba(250,204,21,0.24)", "accent": "#facc15", "accent2": "#94a3b8", "text": "#f3f6ff", "muted": "#93a0bc", "bt": "#fff"}, "grad": ["#facc15", "#94a3b8", "#1e40af"], "font": "Georgia,'Times New Roman',serif", "radius": "12px"},
    "founder": {"name": "Founder", "pal": {"bg": "#0c0c0f", "bg2": "#141418", "card": "rgba(255,255,255,0.06)", "line": "rgba(163,230,53,0.22)", "accent": "#a3e635", "accent2": "#f8fafc", "text": "#f4f6ee", "muted": "#9aa0a8", "bt": "#0c0c0f"}, "grad": ["#a3e635", "#4ade80", "#f8fafc"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "14px"},
    "sales": {"name": "Deal Closer", "pal": {"bg": "#f7f9fc", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(234,88,12,0.15)", "accent": "#ea580c", "accent2": "#1d4ed8", "text": "#202b3d", "muted": "#67748a", "bt": "#fff"}, "grad": ["#ea580c", "#f97316", "#1d4ed8"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "14px"},
    "consulting": {"name": "Advisory", "pal": {"bg": "#f4f6f9", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(51,65,85,0.14)", "accent": "#334155", "accent2": "#0ea5e9", "text": "#101826", "muted": "#64748b", "bt": "#fff"}, "grad": ["#334155", "#0ea5e9", "#6366f1"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "logistics": {"name": "Freight Line", "pal": {"bg": "#0c1220", "bg2": "#121b30", "card": "rgba(255,255,255,0.05)", "line": "rgba(251,191,36,0.26)", "accent": "#fbbf24", "accent2": "#38bdf8", "text": "#eef4fd", "muted": "#93a2b8", "bt": "#1a1405"}, "grad": ["#fbbf24", "#38bdf8", "#334155"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "realestate": {"name": "Property", "pal": {"bg": "#f1f8f7", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(15,118,110,0.16)", "accent": "#0f766e", "accent2": "#c2410c", "text": "#102824", "muted": "#5f7d78", "bt": "#fff"}, "grad": ["#0f766e", "#14b8a6", "#c2410c"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "16px"},
    "hr": {"name": "People Ops", "pal": {"bg": "#faf5ff", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(147,51,234,0.16)", "accent": "#9333ea", "accent2": "#ec4899", "text": "#241033", "muted": "#77628d", "bt": "#fff"}, "grad": ["#9333ea", "#ec4899", "#6366f1"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "20px"},
    "doctor": {"name": "Med Life", "pal": {"bg": "#effbfa", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(13,148,136,0.16)", "accent": "#0d9488", "accent2": "#38bdf8", "text": "#0c2b28", "muted": "#5e807b", "bt": "#fff"}, "grad": ["#0d9488", "#38bdf8", "#14b8a6"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "psychology": {"name": "Mindful", "pal": {"bg": "#f5f3ff", "bg2": "#fdfcff", "card": "rgba(255,255,255,0.85)", "line": "rgba(124,58,237,0.16)", "accent": "#7c3aed", "accent2": "#94a3b8", "text": "#211541", "muted": "#7b7295", "bt": "#fff"}, "grad": ["#7c3aed", "#a78bfa", "#94a3b8"], "font": "Georgia,'Times New Roman',serif", "radius": "20px"},
    "nutrition": {"name": "Vitality", "pal": {"bg": "#f2fbf3", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(22,163,74,0.18)", "accent": "#16a34a", "accent2": "#f59e0b", "text": "#122718", "muted": "#71876f", "bt": "#fff"}, "grad": ["#16a34a", "#84cc16", "#f59e0b"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "20px"},
    "fitness": {"name": "Peak Form", "pal": {"bg": "#0d0d10", "bg2": "#141418", "card": "rgba(255,255,255,0.05)", "line": "rgba(239,68,68,0.28)", "accent": "#ef4444", "accent2": "#f8fafc", "text": "#f6f6f6", "muted": "#9aa0a8", "bt": "#fff"}, "grad": ["#ef4444", "#f97316", "#f8fafc"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "6px"},
    "veterinary": {"name": "Paws & Care", "pal": {"bg": "#fff7ec", "bg2": "#fffdf5", "card": "rgba(255,255,255,0.85)", "line": "rgba(234,88,12,0.16)", "accent": "#ea580c", "accent2": "#a78bfa", "text": "#331c0e", "muted": "#8d7c6a", "bt": "#fff"}, "grad": ["#ea580c", "#f59e0b", "#a78bfa"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "20px"},
    "uiux": {"name": "Interface Lab", "pal": {"bg": "#0d0a1d", "bg2": "#141031", "card": "rgba(255,255,255,0.06)", "line": "rgba(240,171,252,0.26)", "accent": "#f0abfc", "accent2": "#22d3ee", "text": "#f5effa", "muted": "#a294bb", "bt": "#fff"}, "grad": ["#f0abfc", "#22d3ee", "#818cf8"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "26px"},
    "interior": {"name": "Interiors", "pal": {"bg": "#faf6f0", "bg2": "#fffdf9", "card": "rgba(255,255,255,0.85)", "line": "rgba(194,65,12,0.15)", "accent": "#c2410c", "accent2": "#78716c", "text": "#2c2118", "muted": "#8b8272", "bt": "#fff"}, "grad": ["#c2410c", "#a8a29e", "#78716c"], "font": "Georgia,'Times New Roman',serif", "radius": "10px"},
    "fashion": {"name": "Runway", "pal": {"bg": "#120a0c", "bg2": "#190e12", "card": "rgba(255,255,255,0.05)", "line": "rgba(234,179,8,0.28)", "accent": "#eab308", "accent2": "#fdf2f8", "text": "#f7ede9", "muted": "#a08d89", "bt": "#1a1206"}, "grad": ["#eab308", "#f472b6", "#ffffff"], "font": "Georgia,'Times New Roman',serif", "radius": "4px"},
    "videographer": {"name": "Cinema", "pal": {"bg": "#0b0b12", "bg2": "#10101a", "card": "rgba(255,255,255,0.05)", "line": "rgba(185,28,28,0.3)", "accent": "#f43f5e", "accent2": "#f59e0b", "text": "#f3eef4", "muted": "#9a8f9f", "bt": "#fff"}, "grad": ["#b91c1c", "#f59e0b", "#eab308"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "12px"},
    "musician": {"name": "Stage Light", "pal": {"bg": "#0f0718", "bg2": "#170b22", "card": "rgba(255,255,255,0.05)", "line": "rgba(217,70,239,0.3)", "accent": "#d946ef", "accent2": "#fbbf24", "text": "#f7ecfa", "muted": "#a18bab", "bt": "#fff"}, "grad": ["#d946ef", "#fbbf24", "#8b5cf6"], "font": "Georgia,'Times New Roman',serif", "radius": "22px"},
    "artist": {"name": "Gallery White", "pal": {"bg": "#ffffff", "bg2": "#fafafa", "card": "rgba(255,255,255,0.85)", "line": "rgba(17,24,39,0.12)", "accent": "#111827", "accent2": "#f472b6", "text": "#181818", "muted": "#6b7280", "bt": "#fff"}, "grad": ["#111827", "#f472b6", "#6366f1"], "font": "Georgia,'Times New Roman',serif", "radius": "24px"},
    "writer": {"name": "Editorial", "pal": {"bg": "#fbfaf5", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(41,37,36,0.14)", "accent": "#292524", "accent2": "#b45309", "text": "#231f1c", "muted": "#7a736c", "bt": "#fff"}, "grad": ["#292524", "#b45309", "#78716c"], "font": "Georgia,'Times New Roman',serif", "radius": "8px"},
    "journalist": {"name": "Press Room", "pal": {"bg": "#f4f4f5", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(185,28,28,0.16)", "accent": "#b91c1c", "accent2": "#18181b", "text": "#1d1d1f", "muted": "#6b7280", "bt": "#fff"}, "grad": ["#b91c1c", "#18181b", "#52525b"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "6px"},
    "translator": {"name": "Polyglot", "pal": {"bg": "#effcfa", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(13,148,136,0.16)", "accent": "#0d9488", "accent2": "#7c3aed", "text": "#0e2b27", "muted": "#5f7f79", "bt": "#fff"}, "grad": ["#0d9488", "#7c3aed", "#14b8a6"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "18px"},
    "contentcreator": {"name": "Creator Vibes", "pal": {"bg": "#fff1f4", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(225,29,72,0.16)", "accent": "#e11d48", "accent2": "#7c3aed", "text": "#2c1420", "muted": "#8d6f79", "bt": "#fff"}, "grad": ["#e11d48", "#7c3aed", "#f59e0b"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "22px"},
    "architecture": {"name": "Blueprint", "pal": {"bg": "#eef5ff", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(29,78,216,0.18)", "accent": "#1d4ed8", "accent2": "#0f172a", "text": "#101b30", "muted": "#5f7291", "bt": "#fff"}, "grad": ["#1d4ed8", "#38bdf8", "#0f172a"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "4px"},
    "researcher": {"name": "Research Desk", "pal": {"bg": "#0c1224", "bg2": "#111a33", "card": "rgba(255,255,255,0.05)", "line": "rgba(56,189,248,0.26)", "accent": "#38bdf8", "accent2": "#a5b4fc", "text": "#eef4fd", "muted": "#93a2ba", "bt": "#fff"}, "grad": ["#38bdf8", "#a5b4fc", "#6366f1"], "font": "Georgia,'Times New Roman',serif", "radius": "16px"},
    "scientist": {"name": "Quantum", "pal": {"bg": "#050810", "bg2": "#0a0f1c", "card": "rgba(255,255,255,0.05)", "line": "rgba(34,211,238,0.26)", "accent": "#22d3ee", "accent2": "#a78bfa", "text": "#eaf6fc", "muted": "#8fa0b4", "bt": "#fff"}, "grad": ["#22d3ee", "#a78bfa", "#0ea5e9"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "14px"},
    "historian": {"name": "Archive", "pal": {"bg": "#f6f1e6", "bg2": "#fdfaf2", "card": "rgba(255,255,255,0.8)", "line": "rgba(146,64,14,0.16)", "accent": "#92400e", "accent2": "#78716c", "text": "#2b2418", "muted": "#877d6c", "bt": "#fff"}, "grad": ["#92400e", "#b45309", "#78716c"], "font": "Georgia,'Times New Roman',serif", "radius": "6px"},
    "chef": {"name": "Culinary", "pal": {"bg": "#fff6ec", "bg2": "#fffdf8", "card": "rgba(255,255,255,0.85)", "line": "rgba(154,52,18,0.16)", "accent": "#9a3412", "accent2": "#f59e0b", "text": "#2f1707", "muted": "#8f7862", "bt": "#fff"}, "grad": ["#9a3412", "#f59e0b", "#ea580c"], "font": "Georgia,'Times New Roman',serif", "radius": "14px"},
    "hospitality": {"name": "Welcome", "pal": {"bg": "#fbf6ec", "bg2": "#fffdf8", "card": "rgba(255,255,255,0.85)", "line": "rgba(180,83,9,0.16)", "accent": "#b45309", "accent2": "#0f766e", "text": "#2b2113", "muted": "#8b7f6c", "bt": "#fff"}, "grad": ["#b45309", "#0f766e", "#f59e0b"], "font": "Georgia,'Times New Roman',serif", "radius": "12px"},
    "tourism": {"name": "Wanderlust", "pal": {"bg": "#eefaff", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(14,165,233,0.16)", "accent": "#0ea5e9", "accent2": "#f59e0b", "text": "#0c2434", "muted": "#5f7f92", "bt": "#fff"}, "grad": ["#0ea5e9", "#14b8a6", "#f59e0b"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "22px"},
    "trades": {"name": "Tradesman", "pal": {"bg": "#0e0c0a", "bg2": "#151210", "card": "rgba(255,255,255,0.05)", "line": "rgba(249,115,22,0.28)", "accent": "#f97316", "accent2": "#facc15", "text": "#f6efe6", "muted": "#a0907e", "bt": "#fff"}, "grad": ["#f97316", "#facc15", "#78716c"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "8px"},
    "beauty": {"name": "Glow Studio", "pal": {"bg": "#fdf2f7", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(219,39,119,0.16)", "accent": "#db2777", "accent2": "#f472b6", "text": "#331221", "muted": "#8d6b7c", "bt": "#fff"}, "grad": ["#db2777", "#f472b6", "#fb923c"], "font": "Georgia,'Times New Roman',serif", "radius": "26px"},
    "aviation": {"name": "Aviation", "pal": {"bg": "#f0f9ff", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(3,105,161,0.16)", "accent": "#0369a1", "accent2": "#e2e8f0", "text": "#0c2536", "muted": "#5f7d92", "bt": "#fff"}, "grad": ["#0369a1", "#38bdf8", "#94a3b8"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "14px"},
    "agriculture": {"name": "Harvest", "pal": {"bg": "#f7fee7", "bg2": "#fdfff8", "card": "rgba(255,255,255,0.85)", "line": "rgba(101,163,13,0.2)", "accent": "#65a30d", "accent2": "#92400e", "text": "#1c2b0e", "muted": "#7c8a68", "bt": "#fff"}, "grad": ["#65a30d", "#84cc16", "#92400e"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "16px"},
    "social": {"name": "Compassion", "pal": {"bg": "#fff4ec", "bg2": "#fffdf9", "card": "rgba(255,255,255,0.85)", "line": "rgba(234,88,12,0.16)", "accent": "#ea580c", "accent2": "#7c3aed", "text": "#2f1a0e", "muted": "#8d7a68", "bt": "#fff"}, "grad": ["#ea580c", "#ec4899", "#f59e0b"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "22px"},
    "government": {"name": "Civic", "pal": {"bg": "#f5f7fa", "bg2": "#ffffff", "card": "rgba(255,255,255,0.85)", "line": "rgba(30,58,138,0.16)", "accent": "#1e3a8a", "accent2": "#dc2626", "text": "#101b33", "muted": "#64748b", "bt": "#fff"}, "grad": ["#1e3a8a", "#3b82f6", "#dc2626"], "font": "'Inter','Segoe UI',system-ui,-apple-system,sans-serif", "radius": "10px"},
    "legal": {"name": "Justice", "pal": {"bg": "#0b0d12", "bg2": "#12151d", "card": "rgba(255,255,255,0.05)", "line": "rgba(234,179,8,0.26)", "accent": "#eab308", "accent2": "#94a3b8", "text": "#f2f3f7", "muted": "#98a0ad", "bt": "#1a1405"}, "grad": ["#eab308", "#94a3b8", "#52525b"], "font": "Georgia,'Times New Roman',serif", "radius": "6px"},
}


# ---------------------------------------------------------------------------
# PORTFOLIO RENDERER (ported from cvforge/portfolio.py, self-contained)
# ---------------------------------------------------------------------------
def _esc(v) -> str:
    return str(v if v is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _initials(name: str) -> str:
    parts = [p for p in (name or "CV").replace("·", " ").split() if p]
    return "".join(p[0].upper() for p in parts[:2]) or "CV"


def generate_portfolio_html(cv: dict, theme_key: str | None = None, language: str = "en", seed: int | None = None) -> str:
    """Procedural design: (theme palette + seed) -> unique structured design."""
    t = THEMES.get(theme_key or cv.get("domain") or "generic", THEMES["generic"])
    return render_portfolio_html(cv, t, language, seed)


"""CVForge — INFINITE DESIGN ENGINE (procedural portfolio designs).

Every (seed, theme) pair produces a different STRUCTURAL design: layout
archetype, background treatment, typography, shape language, card style,
section accents. Same seed + same theme = identical design (shareable);
new seed = brand-new design. Deterministic (mulberry32) + identical in
Python and JS (landing page/offline playground use the JS port).
"""

import html as _html
import json as _json

# ---------------------------------------------------------------- RNG ----
def _h32(s: str) -> int:
    h = 0
    for ch in s:
        h = (h * 31 + ord(ch)) & 0xFFFFFFFF
    return h


def _mulberry32(seed: int):
    s = seed & 0xFFFFFFFF

    def rnd():
        nonlocal s
        s = (s + 0x6D2B79F5) & 0xFFFFFFFF
        t = s
        t = ((t ^ (t >> 15)) * (t | 1)) & 0xFFFFFFFF
        t = (t ^ (t + (((t ^ (t >> 7)) * (t | 61)) & 0xFFFFFFFF))) & 0xFFFFFFFF
        return ((t ^ (t >> 14)) & 0xFFFFFFFF) / 4294967296

    return rnd


# ------------------------------------------------------------ vocab ------
LAYOUTS = ["split", "hero", "bento", "editorial", "terminal", "glass", "white", "magazine"]
BGS = ["blobs", "mesh", "grid", "dots", "scan", "waves", "stripes", "plain"]
FONTS = ["display", "serif", "mono", "sans", "round"]
RADII = ["0px", "10px", "22px", "32px"]
CHIPS = ["999px", "8px", "22px", "6px"]
CARDS = ["fill", "outline", "glass", "hard"]
ACCENTS = ["bar", "underline", "num", "dot"]
HEROS = ["avatar", "initials", "split", "min"]
FONT_STACKS = {
    "display": "'Barlow Condensed','Arial Narrow','Inter',system-ui,sans-serif",
    "serif": "'Playfair Display',Georgia,'Times New Roman',serif",
    "mono": "'JetBrains Mono','Fira Code',ui-monospace,Menlo,Consolas,monospace",
    "sans": "'Inter','Segoe UI',system-ui,sans-serif",
    "round": "'Nunito','Segoe UI',system-ui,sans-serif",
}


def design_spec(theme_key: str | None, seed: int | None = None) -> dict:
    seed = _h32(theme_key or "cvforge") if seed is None else (int(seed) & 0xFFFFFFFF)
    r = _mulberry32(seed or 1)
    i = lambda n: int(r() * n)  # noqa: E731
    return {
        "seed": seed,
        "layout": LAYOUTS[i(len(LAYOUTS))],
        "bg": BGS[i(len(BGS))],
        "font": FONTS[i(len(FONTS))],
        "shape": i(len(RADII)),
        "chip": i(len(CHIPS)),
        "card": CARDS[i(len(CARDS))],
        "accent": ACCENTS[i(len(ACCENTS))],
        "hero": HEROS[i(len(HEROS))],
        "motion": r() < 0.85,
        "grain": r() < 0.35,
    }


def _des_esc(v) -> str:
    return _html.escape(str(v or ""))


def _des_ini(name: str) -> str:
    parts = [p for p in (name or "CV").replace("·", " ").split() if p]
    return "".join(p[0].upper() for p in parts[:2]) if parts else "CV"


# ------------------------------------------------------------ render -----
def render_portfolio_html(cv: dict, theme: dict, language: str = "en", seed: int | None = None) -> str:
    pal = dict(theme.get("pal", {}))
    g = theme.get("grad", ["#6366f1", "#22d3ee", "#a855f7"])
    spec = design_spec(theme.get("name", "cvforge") or "cvforge", seed)
    p = dict(pal)
    p.setdefault("line", "rgba(255,255,255,0.09)")
    p.setdefault("bt", "#fff")
    p.setdefault("bg2", p.get("bg", "#0a0c14"))
    bg_hex = p.get("bg", "#000000")
    dark = (bg_hex.startswith("#") and len(bg_hex) >= 7 and int(bg_hex[1:3], 16) < 200) or spec["layout"] == "terminal"

    # layout-forced bases
    if spec["layout"] == "terminal":
        p.update(bg="#0b0e13", bg2="#10141c", card="rgba(0,255,65,0.04)",
                 line="rgba(255,255,255,0.12)", text="#e9f5ec", muted="#7f9f8a", bt="#031408")
        dark = True
    if spec["layout"] == "white":
        p.update(bg="#ffffff", bg2="#f7f8fa", card="#ffffff", line="rgba(15,23,42,0.12)",
                 text="#0f172a", muted="#64748b", bt="#0f172a")
        dark = False

    # content -------------------------------------------------------------
    name = _des_esc(cv.get("name") or "Your Name")
    title = _des_esc(cv.get("title") or cv.get("domain_label") or "Professional")
    label = _des_esc(cv.get("domain_label") or "Professional")
    summary = _des_esc((cv.get("summary") or "").strip() or f"{name} — {label} committed to quality, impact, and continuous growth.")
    email = _des_esc(cv.get("email") or "")
    website = _des_esc(cv.get("website") or "")
    phone = _des_esc(cv.get("phone") or "")
    loc = _des_esc(cv.get("location") or "")
    href = (f"mailto:{email}" if email else
            (website if website.startswith("http") else f"https://{website}" if website else "#"))
    ini = _des_ini(cv.get("name"))
    skills = "".join(f'<span class="chip">{_des_esc(s)}</span>' for s in (cv.get("skills") or [])[:24]) or '<span class="chip">Core skills</span>'
    langs = "".join(f'<span class="chip ghost">{_des_esc(l)}</span>' for l in (cv.get("languages") or [])[:10])

    def tl(items):
        if not items:
            return ""
        return '<div class="timeline">' + "".join(
            '<div class="tl-item"><div class="tl-dot"></div><div class="tl-body">'
            f'<h3>{_des_esc(it.get("title") or it.get("role") or it.get("degree") or "")}</h3>'
            f'<p class="tl-meta">{_des_esc(it.get("period") or it.get("company") or it.get("school") or "")}</p>'
            + "".join(f"<p>{_des_esc(d)}</p>" for d in it.get("details", []) if d and d.strip())
            + "</div></div>" for it in items) + "</div>"

    exp = tl(cv.get("experience") or [])
    edu = tl(cv.get("education") or [])
    projs = "".join(
        '<article class="card cproj"><h3>' + _des_esc(x.get("title") or "") + "</h3>"
        + "".join(f"<p>{_des_esc(d)}</p>" for d in x.get("details", []) if d and d.strip()) + "</article>"
        for x in (cv.get("projects") or [])[:9]) or '<p class="muted">Contact for a walkthrough of selected work.</p>'
    words = [title, label, "Available for new opportunities"]

    def head(label_txt, idx):
        a = spec["accent"]
        if a == "bar":
            return f'<h2><span class="bar"></span>{label_txt}</h2>'
        if a == "underline":
            return f'<h2 class="ul">{label_txt}</h2>'
        if a == "num":
            return f'<h2 class="num"><b>{idx:02d}</b>{label_txt}</h2>'
        return f'<h2 class="dotac"><i></i>{label_txt}</h2>'

    sec_about = f'<section id="about" class="reveal">{head("About", 1)}<p class="sub">{label} · {summary}</p></section>'
    sec_skills = f'<section id="skills" class="reveal">{head("Skills", 2)}<div class="chips">{skills}</div>' + (f'<div class="chips" style="margin-top:14px">{langs}</div>' if langs else "") + "</section>"
    sec_exp = f'<section id="experience" class="reveal">{head("Experience", 3)}' + (exp or '<p class="sub">Detailed experience available on request.</p>') + "</section>"
    sec_edu = f'<section id="education" class="reveal">{head("Education", 4)}{edu}</section>' if edu else ""
    sec_proj = f'<section id="projects" class="reveal">{head("Projects", 5)}<div class="gridP">{projs}</div></section>'
    footer = ('<footer id="contact" class="reveal"><div class="big"><a class="big-a" href="' + href + '">'
              + (email or website or "Let’s work together") + "</a></div>"
              + (f'<div class="meta">📍 {loc} · ☎ {phone}</div>' if (loc or phone) else "")
              + f'<p class="tiny">Generated by <b>CVForge</b> · Design #{spec["seed"]} · {name} · {label}</p></footer>')

    hero_map = {
        "avatar": f'<header class="hero reveal"><div class="avatar">{ini}</div><h1>{name}<br><span class="g">{title}</span></h1>'
                  f'<div class="typed">{title}</div><p class="lead">{summary}</p><div class="cta-row">'
                  f'<a class="btn primary" href="{href}">Get in touch</a><a class="btn ghost" href="#projects">See work ↓</a></div></header>',
        "initials": f'<header class="hero reveal"><div class="avatar">{ini}</div><h1>{name}<br><span class="g">{title}</span></h1>'
                    f'<div class="typed">{title}</div><p class="lead">{summary}</p><div class="cta-row">'
                    f'<a class="btn primary" href="{href}">Get in touch</a><a class="btn ghost" href="#projects">See work ↓</a></div></header>',
        "split": f'<header class="hero reveal"><div class="avatar round">{ini}</div><div>'
                 f'<h1>{name}<br><span class="g">{title}</span></h1><div class="typed">{title}</div>'
                 f'<p class="lead">{summary}</p><div class="cta-row"><a class="btn primary" href="{href}">Get in touch</a>'
                 f'<a class="btn ghost" href="#projects">See work ↓</a></div></div></header>',
        "min": f'<header class="hero reveal"><h1>{name} <span class="g">— {title}</span></h1>'
               f'<p class="lead">{summary}</p><div class="cta-row"><a class="btn primary" href="{href}">Get in touch</a>'
               f'<a class="btn ghost" href="#projects">See work ↓</a></div></header>',
    }
    hero = hero_map[spec["hero"]]

    nav = ('<nav><div class="wrap"><a class="logo" href="#top">' + ini + "</a><div class=\"nav-links\">"
           '<a href="#about">About</a><a href="#skills">Skills</a><a href="#experience">Experience</a>'
           '<a href="#projects">Projects</a><a href="#contact">Contact</a></div></div></nav>')

    L = spec["layout"]
    if L == "split":
        side = ('<aside class="side-card"><div class="avatar">' + ini + '</div><h3 class="side-name">' + name
                + '</h3><p class="muted">' + title + '</p><div class="contact-mini">'
                + (f'<a href="mailto:{email}">✉ {email}</a>' if email else "")
                + (f'<span>☎ {phone}</span>' if phone else "")
                + (f'<span>📍 {loc}</span>' if loc else "") + "</div></aside>")
        inner = ('<div class="shell"><div class="side-col">' + side + '</div><main class="main-col">'
                 + hero + sec_about + sec_skills + sec_exp + sec_edu + sec_proj + footer + "</main></div>")
    elif L == "bento":
        inner = ('<div class="shell">'
                 f'<section class="bcard b-about reveal">{head("About", 1)}<p class="sub">{label} · {summary}</p></section>'
                 f'<section class="bcard b-skills reveal">{head("Skills", 2)}<div class="chips">{skills}</div></section>'
                 f'<section class="bcard b-exp reveal">{head("Experience", 3)}{exp or "<p class=\"sub\">On request.</p>"}</section>'
                 + (f'<section class="bcard b-edu reveal">{head("Education", 4)}{edu}</section>' if edu else "")
                 + f'<section class="bcard b-proj reveal">{head("Projects", 5)}<div class="gridP">{projs}</div></section>'
                 + '<footer class="bcard b-foot" id="contact"><div class="big"><a class="big-a" href="' + href + '">'
                 + (email or website or "Let’s work together") + '</a></div><p class="tiny">Generated by <b>CVForge</b> · Design #' + str(spec["seed"]) + '</p></footer>'
                 + "</div>")
    elif L == "magazine":
        band = ('<div class="mband"><div class="wrap"><div class="mband-left"><div class="mma">PORTFOLIO</div>'
                f'<h1 class="mh1">{name}</h1><p class="msub">{title}</p></div>'
                f'<div class="mband-right"><div class="avatar big">{ini}</div>'
                + (f'<div class="mm">{loc}</div>' if loc else "") + "</div></div></div>")
        inner = ('<div class="shell-mz">' + band + '<main class="wrap" id="top">' + sec_about + sec_skills
                 + sec_exp + sec_edu + sec_proj + footer.replace('class="big"', 'class="big"') + "</main></div>")
    else:
        inner = ('<main class="wrap" id="top">' + hero + sec_about + sec_skills + sec_exp + sec_edu
                 + sec_proj + footer + "</main>")

    # css present for editorial wrapper etc.
    if L == "editorial":
        inner = inner.replace('<main class="wrap" id="top">', '<main class="container-max" id="top">')

    # ------------------------------------------------------- css assembly
    grad = f"linear-gradient(135deg,{g[0]},{g[1]},{g[2]})"
    grad_len = f"linear-gradient(90deg,{g[0]},{g[1]},{g[2]})"
    fbody = FONT_STACKS["sans"] if spec["font"] in ("sans", "round", "display") else FONT_STACKS["serif"] if spec["font"] == "serif" else FONT_STACKS["mono"]
    pal_vars = (f"--bg:{p['bg']};--bg2:{p['bg2']};--card:{p['card']};--line:{p['line']};"
                f"--accent:{p['accent']};--accent2:{p['accent2']};--text:{p['text']};--muted:{p['muted']};--bt:{p['bt']};"
                f"--grad:{grad};--grad-len:{grad_len};--r:{RADII[spec['shape']]};--chipr:{CHIPS[spec['chip']]};"
                f"--fbody:{fbody};--fhead:{FONT_STACKS[spec['font']]};")
    css = (_CSS_TMPL
           + _LAYOUT_CSS[L] + _BG_CSS[spec["bg"]] + _FONT_CSS[spec["font"]]
           + _SHAPE_CSS[spec["shape"]] + _CARD_CSS[spec["card"]] + _HERO_CSS[spec["hero"]])
    css = (css.replace("@@PAL@@", pal_vars).replace("@@ACCENT@@", g[0])
              .replace("@@ACCENT2@@", g[1]).replace("@@C3@@", g[2] if len(g) > 2 else g[1]))

    rtl = (language or "en").lower().startswith("ar")
    langattr, dirattr = ("ar", ' dir="rtl"') if rtl else ("en", "")
    typed = ""
    if spec["motion"]:
        w = _json.dumps(words, ensure_ascii=False)
        typed = ("<script>(function(){var w=" + w
                 + ";var el=document.querySelector('.typed');if(!el)return;var wi=0,ci=0,del=false;"
                 "function t(){var s=w[wi]||'';el.textContent=s.slice(0,ci)+'▌';"
                 "if(!del&&ci<s.length){ci++;setTimeout(t,55);}else if(!del){del=true;setTimeout(t,1500);}"
                 "else if(ci>0){ci--;setTimeout(t,24);}else{del=false;wi=(wi+1)%w.length;setTimeout(t,300);}}t();"
                 "var io=new IntersectionObserver(function(es){es.forEach(function(e){if(e.isIntersecting)"
                 "{e.target.classList.add('in');io.unobserve(e.target);}});},{threshold:.12});"
                 "document.querySelectorAll('.reveal').forEach(function(e){io.observe(e);});})();</script>")
    grain = '<div class="grain"></div>' if spec["grain"] else ""

    return ('<!DOCTYPE html><html lang="' + langattr + '"' + dirattr + '><head><meta charset="UTF-8">'
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">'
            f'<title>{name} — {title}</title><style>{css}</style></head>'
            f'<body class="ls-{L} {"dark" if dark else "light"} bgbg-{spec["bg"]}">'
            f'<div class="bg" aria-hidden="true">{_BG_HTML[spec["bg"]]}</div>{grain}{nav}{inner}{typed}</body></html>')


# ----------------------------------------------------------------- CSS ---
_CSS_TMPL = """:root{@@PAL@@}
*{margin:0;padding:0;box-sizing:border-box}
html{scroll-behavior:smooth}
body{background:var(--bg);color:var(--text);font-family:var(--fbody);line-height:1.65;overflow-x:hidden;position:relative;min-height:100vh}
a{color:inherit;text-decoration:none}
::selection{background:var(--accent);color:var(--bt)}
.wrap{max-width:1040px;margin:0 auto;padding:0 22px}
.container-max{max-width:860px;margin:0 auto;padding:0 22px}
.muted{color:var(--muted)}
.grain{position:fixed;inset:0;z-index:90;pointer-events:none;opacity:.05;background-image:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='160' height='160'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.8' numOctaves='2'/%3E%3C/filter%3E%3Crect width='160' height='160' filter='url(%23n)' opacity='0.6'/%3E%3C/svg%3E")}
.bg{position:fixed;inset:0;z-index:-2;overflow:hidden}
nav{position:sticky;top:0;z-index:50;background:color-mix(in srgb,var(--bg) 92%,transparent);backdrop-filter:blur(12px);border-bottom:1px solid var(--line)}
nav .wrap{display:flex;align-items:center;justify-content:space-between;padding:14px 22px}
.logo{font-weight:900;font-size:1.1rem;letter-spacing:.5px;background:var(--grad-len);-webkit-background-clip:text;background-clip:text;color:transparent}
.nav-links{display:flex;gap:4px;flex-wrap:wrap}
.nav-links a{font-size:.82rem;padding:7px 13px;border-radius:999px;color:var(--muted);font-weight:600}
.nav-links a:hover{color:var(--text);background:var(--card)}
.avatar{width:92px;height:92px;border-radius:26px;background:var(--grad);display:grid;place-items:center;font-weight:900;font-size:1.9rem;color:var(--bt);box-shadow:0 16px 44px @@ACCENT@@44}
.hero{margin-top:34px;padding:46px 0 40px}
.hero h1{font-family:var(--fhead);font-size:clamp(2.4rem,7vw,4.4rem);line-height:1.02;letter-spacing:-.02em;font-weight:800}
.hero h1 .g{background:var(--grad-len);-webkit-background-clip:text;background-clip:text;color:transparent}
.typed{font-family:var(--fhead);font-size:clamp(1rem,2.6vw,1.3rem);color:var(--accent2);min-height:1.6em;font-weight:700}
.lead{color:var(--muted);margin-top:14px;max-width:660px}
.cta-row{display:flex;gap:12px;margin-top:26px;flex-wrap:wrap}
.btn{padding:13px 26px;font-weight:800;font-size:.92rem;border:1px solid var(--line);border-radius:999px;display:inline-block;transition:.25s}
.btn.primary{background:var(--grad);color:var(--bt);box-shadow:0 10px 30px @@ACCENT@@44}
.btn.primary:hover{transform:translateY(-3px)}
.btn.ghost:hover{background:var(--card)}
section{margin-top:52px}
h2{font-family:var(--fhead);font-size:clamp(1.4rem,4vw,2rem);margin-bottom:14px;letter-spacing:-.01em}
h2 .bar{display:inline-block;width:44px;height:5px;border-radius:99px;background:var(--grad-len);margin-inline-start:8px;vertical-align:middle}
h2.ul{border-bottom:2px solid var(--line);padding-bottom:8px}
h2.num{border:1px solid var(--line);padding:10px 16px;border-radius:var(--r);background:var(--card);font-weight:900}
h2.num b{color:var(--accent);margin-inline-end:10px;font-family:var(--fbody)}
h2.dotac{display:flex;align-items:center;gap:10px}
h2.dotac i{width:12px;height:12px;border-radius:50%;background:var(--grad);box-shadow:0 0 12px @@ACCENT@@}
.sub{color:var(--muted);margin-bottom:22px}
.chips{display:flex;flex-wrap:wrap;gap:9px}
.chip{padding:8px 16px;border-radius:var(--chipr);font-size:.84rem;background:var(--card);border:1px solid var(--line);font-weight:600}
.chip.ghost{background:transparent}
.timeline{position:relative;padding-inline-start:26px}
.timeline::before{content:"";position:absolute;inset-block:4px;inset-inline-start:9px;width:2px;background:linear-gradient(var(--accent),transparent);opacity:.5}
.tl-item{position:relative;margin-bottom:18px}
.tl-dot{position:absolute;inset-inline-start:-26px;top:2px;width:18px;height:18px;border-radius:50%;background:var(--bg);border:2px solid var(--accent);display:grid;place-items:center}
.tl-body{padding:16px 18px;background:var(--card);border:1px solid var(--line);border-radius:var(--r)}
.tl-body h3{font-size:1.02rem}
.tl-meta{color:var(--accent2);font-size:.8rem;font-weight:700;margin-top:2px}
.tl-body p{color:var(--muted);font-size:.88rem;margin-top:5px}
.gridP{display:grid;grid-template-columns:repeat(auto-fit,minmax(250px,1fr));gap:16px}
.cproj{padding:20px;transition:.25s}
.cproj:hover{transform:translateY(-5px);border-color:@@ACCENT@@66}
.cproj h3{font-size:1.04rem;margin-bottom:8px}
.cproj p{color:var(--muted);font-size:.88rem}
.card{background:var(--card);border:1px solid var(--line);border-radius:var(--r)}
footer{margin:60px 0 34px;text-align:center;padding:30px;background:var(--card);border:1px solid var(--line);border-radius:var(--r)}
.big{font-size:clamp(1.2rem,4vw,1.8rem);font-weight:900}
.big-a{background:var(--grad-len);-webkit-background-clip:text;background-clip:text;color:transparent}
.meta{color:var(--muted);margin-top:8px;font-size:.88rem}
.tiny{color:var(--muted);font-size:.72rem;margin-top:16px;opacity:.75}
.contact-mini{display:flex;flex-direction:column;gap:7px;margin-top:16px;font-size:.84rem;color:var(--muted);word-break:break-word}
.contact-mini a:hover{color:var(--accent2)}
.reveal{opacity:0;transform:translateY(22px);transition:opacity .6s ease,transform .6s ease}
.reveal.in{opacity:1;transform:none}
@media(max-width:700px){.shell{grid-template-columns:1fr!important}.side-col{display:none}.mband .wrap{grid-template-columns:1fr}.wrap,.container-max{padding:0 16px}}
@media (prefers-reduced-motion:reduce){*{animation:none!important;transition:none!important}.reveal{opacity:1;transform:none}}
"""

_LAYOUT_CSS = {
 "split": ".shell{display:grid;grid-template-columns:320px 1fr;gap:30px;max-width:1120px;margin:0 auto;padding:0 22px}.side-col{padding-top:30px}.side-card{position:sticky;top:82px;background:var(--card);border:1px solid var(--line);border-radius:var(--r);padding:26px;text-align:center}.side-card .avatar{margin:0 auto 18px}.side-name{font-family:var(--fhead);font-size:1.5rem}.main-col{padding-top:30px;min-width:0}",
 "hero": ".hero{text-align:center;padding:72px 0 48px}.hero .avatar{margin:0 auto 20px}.hero h1{font-size:clamp(2.8rem,9vw,5.6rem)}.lead{margin:16px auto 0}.cta-row{justify-content:center}section{max-width:820px;margin-inline:auto}",
 "bento": ".shell{max-width:1080px;margin:0 auto;padding:0 22px;display:grid;grid-template-columns:repeat(12,1fr);gap:16px}.bcard{background:var(--card);border:1px solid var(--line);border-radius:var(--r);padding:24px;grid-column:span 6;margin-top:52px}.b-about{grid-column:span 12;margin-top:34px;background:var(--grad);color:var(--bt);border:none}.b-about h2{color:var(--bt)}.b-about h2 .bar{background:#fff}.b-about .sub{color:color-mix(in srgb,var(--bt) 78%,transparent)}.b-skills{grid-column:span 5}.b-exp{grid-column:span 7}.b-edu{grid-column:span 5}.b-proj{grid-column:span 7}.b-foot{grid-column:span 12}.b-about h2.num,.b-about h2.dotac i{color:var(--bt)}@media(max-width:800px){.bcard{grid-column:span 12!important}}",
 "editorial": ".hero{padding:60px 0 30px}.hero h1{font-family:var(--fhead);font-size:clamp(2.6rem,8vw,4.6rem)}section{border-top:1px solid var(--line);padding-top:26px}",
 "terminal": "nav{background:rgba(11,14,19,.92)}.hero{border:1px solid var(--line);border-radius:var(--r);padding:34px;background:rgba(255,255,255,0.02);font-family:var(--fbody)}.hero h1{text-transform:uppercase;letter-spacing:.04em;font-size:clamp(2rem,6vw,3.2rem)}.typed{color:var(--accent)}.chip,.tl-body,.cproj,footer{border-radius:4px}.term{display:none}",
 "glass": ".hero .avatar{border-radius:50%}.tl-body,.cproj,footer{background:color-mix(in srgb,var(--card) 60%,transparent);backdrop-filter:blur(18px)}",
 "white": "body{background:#fff}nav{background:rgba(255,255,255,.9)}h2{color:#0f172a}.lead{color:#64748b}.avatar{box-shadow:0 16px 40px rgba(15,23,42,.14)}",
 "magazine": ".shell-mz .mband{background:var(--grad);color:var(--bt);padding:56px 0 46px}.mband .wrap{display:grid;grid-template-columns:1fr auto;gap:26px;align-items:center}.mband-left .mma{font-size:.72rem;letter-spacing:4px;font-weight:900;opacity:.85}.mh1{font-family:var(--fhead);font-size:clamp(2.4rem,7vw,4.2rem);line-height:1;color:var(--bt)}.msub{color:color-mix(in srgb,var(--bt) 80%,transparent);font-weight:700;margin-top:8px}.mband-right .mm{font-size:.85rem;color:color-mix(in srgb,var(--bt) 80%,transparent);margin-top:10px;text-align:end}.mband .avatar{border-radius:50%;border:3px solid color-mix(in srgb,var(--bt) 60%,transparent)}#top{margin-top:6px}",
}
_BG_CSS = {
 "blobs": ".bgbg-blobs .blob{position:absolute;border-radius:50%;filter:blur(90px);opacity:.3;animation:drift 22s ease-in-out infinite alternate}.bgbg-blobs .b1{width:42vmax;height:42vmax;background:@@ACCENT@@;top:-14vmax;left:-10vmax}.bgbg-blobs .b2{width:34vmax;height:34vmax;background:@@ACCENT2@@;bottom:-12vmax;right:-8vmax;animation-delay:-8s}.bgbg-blobs .b3{width:20vmax;height:20vmax;background:@@C3@@;top:38%;left:55%;animation-delay:-15s;opacity:.2}@keyframes drift{to{transform:translate(6vmax,4vmax) scale(1.12)}}",
 "mesh": ".bgbg-mesh .bg{background:radial-gradient(at 18% 12%,@@ACCENT@@33,transparent 46%),radial-gradient(at 85% 18%,@@ACCENT2@@2e,transparent 42%),radial-gradient(at 70% 85%,@@ACCENT@@26,transparent 45%)}",
 "grid": ".bgbg-grid .bg{background-image:linear-gradient(var(--line) 1px,transparent 1px),linear-gradient(90deg,var(--line) 1px,transparent 1px);background-size:44px 44px;opacity:.55}",
 "dots": ".bgbg-dots .bg{background-image:radial-gradient(var(--accent) 1.2px,transparent 1.3px);background-size:26px 26px;opacity:.22}",
 "scan": ".bgbg-scan .bg{background-image:repeating-linear-gradient(0deg,rgba(255,255,255,0.025) 0 1px,transparent 1px 4px)}.bgbg-scan.light .bg{background-image:repeating-linear-gradient(0deg,rgba(15,23,42,.05) 0 1px,transparent 1px 4px)}",
 "waves": ".bgbg-waves .bg{background:repeating-radial-gradient(circle at 110% -10%,transparent 0 48px,@@ACCENT@@14 48px 50px),repeating-radial-gradient(circle at -10% 110%,transparent 0 60px,@@ACCENT2@@12 60px 62px)}",
 "stripes": ".bgbg-stripes .bg{background-image:repeating-linear-gradient(45deg,transparent 0 26px,@@ACCENT@@0d 26px 28px)}",
 "plain": "",
}
_BG_HTML = {
 "blobs": '<div class="blob b1"></div><div class="blob b2"></div><div class="blob b3"></div>',
 "mesh": "", "grid": "", "dots": "", "scan": "", "waves": "", "stripes": "", "plain": "",
}
_FONT_CSS = {
 "display": "h1{text-transform:uppercase;letter-spacing:.5px}h2{letter-spacing:.02em}",
 "serif": ".hero h1 span,.big-a,.mh1{font-style:italic}",
 "mono": "h1,h2{letter-spacing:.02em}",
 "sans": "",
 "round": "body{letter-spacing:.1px}h2{font-weight:800}",
}
_SHAPE_CSS = ["", "", "", ""]
_SHAPE_CSS[0] = ".avatar{border-radius:50%}.chip{border-radius:999px}"
_SHAPE_CSS[1] = ".avatar{border-radius:12px}"
_SHAPE_CSS[2] = ".avatar{border-radius:30px}"
_SHAPE_CSS[3] = ".avatar{border-radius:6px}"
_CARD_CSS = {
 "fill": "",
 "outline": ".tl-body,.cproj,footer{background:transparent;border:1px solid var(--line)}",
 "glass": ".tl-body,.cproj,footer{background:color-mix(in srgb,var(--card) 55%,transparent);backdrop-filter:blur(16px)}",
 "hard": ".tl-body,.cproj,footer{box-shadow:6px 6px 0 var(--accent)}",
}
_HERO_CSS = {
 "avatar": "",
 "initials": ".hero .avatar{width:120px;height:120px;border-radius:34px;font-size:2.6rem}",
 "split": ".hero{display:grid;grid-template-columns:1fr auto;gap:30px;align-items:center}.hero .avatar{grid-column:2;grid-row:1/3;width:130px;height:130px}",
 "min": ".hero{padding:40px 0 26px}.hero .lead{max-width:520px}.cta-row{margin-top:18px}",
}


# FASTAPI APP — single catch-all dispatcher (no FastAPI injection pitfalls,
# works with ANY path Vercel forwards: /api/health, /health, /api/index, /)
# ---------------------------------------------------------------------------
app = FastAPI(title="CVForge API", version="1.1.0")

_OUT = Path("/tmp/cvforge_api")
_OUT.mkdir(parents=True, exist_ok=True)
_LANDING = Path(__file__).parent / "landing.html"


def _health():
    return {"ok": True, "name": "cvforge-api", "parts": ["cvforge"]}


def _themes():
    return {"ok": True, "count": len(THEMES), "themes": {k: v["name"] for k, v in THEMES.items()}}


def _parse(text: str):
    try:
        return {"ok": True, "cv": parse_cv(text)}
    except Exception as e:  # noqa: BLE001
        return {"ok": False, "error": str(e)}


def _generate(text: str, theme: str | None, language: str, seed: int | None = None):
    if not text.strip():
        return {"ok": False, "error": "no CV text"}
    try:
        cv = parse_cv(text)
        t = THEMES.get(theme or cv.get("domain") or "generic", THEMES["generic"])
        html_content = render_portfolio_html(cv, t, language, seed)
        sp = design_spec(t["name"], seed)
        return {"ok": True, "html": html_content, "chars": len(html_content),
                "domain": cv.get("domain", "generic"), "domain_label": cv.get("domain_label"),
                "design": {"seed": sp["seed"], "layout": sp["layout"], "bg": sp["bg"],
                           "font": sp["font"], "accent": sp["accent"], "hero": sp["hero"]}}
    except Exception as e:  # noqa: BLE001
        return {"ok": False, "error": str(e)}


def _parse_file(path: Path):
    """Extract text from an uploaded CV file (PDF/DOCX/MD/TXT) then parse it."""
    try:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            from pypdf import PdfReader  # lazy: only needed for uploaded PDFs
            reader = PdfReader(str(path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif suffix in (".docx", ".doc"):
            import docx  # lazy: only needed for uploaded Word files
            d = docx.Document(str(path))
            text = "\n".join(p.text for p in d.paragraphs)
        else:  # .md, .txt, or anything else read as text
            text = path.read_text(encoding="utf-8", errors="ignore")
        return _parse(text)
    except Exception as e:  # noqa: BLE001
        return {"ok": False, "error": f"could not read {path.suffix}: {e}"}
    finally:
        path.unlink(missing_ok=True)




def _landing():
    try:
        return Response(content=_LANDING.read_text(encoding="utf-8"), media_type="text/html")
    except Exception:
        return PlainTextResponse("CVForge API — see /api/health", media_type="text/plain")


async def _dispatch(full_path: str, request: Request):
    """Catch-all: whatever path Vercel forwards, dispatch manually."""
    orig = full_path.strip("/")
    path = orig

    # normalize: strip leading api/index, api, or index prefixes
    for prefix in ("api/index", "api", "index"):
        if path == prefix or path.startswith(prefix + "/"):
            path = path[len(prefix):].strip("/")
            break

    if request.method == "GET":
        if path == "health":
            return Response(content=json.dumps(_health()), media_type="application/json")
        if path == "cv/themes":
            return Response(content=json.dumps(_themes()), media_type="application/json")
        if orig in ("", "index.html", "playground", "offline", "demo"):
            return _landing()
        if path == "":  # came from /api or /api/index
            return Response(content=json.dumps(_health()), media_type="application/json")
        return JSONResponse({"ok": False, "error": f"unknown route: /{orig}"}, status_code=404)

    if request.method == "POST":
        # read body manually (JSON or multipart/file)
        ct = request.headers.get("content-type", "")
        text = ""
        theme = None
        language = "en"
        seed = None
        uploaded = None
        try:
            if "multipart" in ct or "application/x-www-form-urlencoded" in ct:
                form = await request.form()
                f = form.get("file")
                if f is not None and hasattr(f, "filename"):
                    tmp = _OUT / f"upload_{uuid.uuid4().hex[:8]}{Path(f.filename or 'cv.txt').suffix.lower()}"
                    tmp.write_bytes(await f.read())
                    uploaded = tmp
                else:
                    text = str(form.get("text", "") or "")
            else:
                raw = await request.body()
                data = json.loads(raw.decode("utf-8") or "{}") if raw else {}
                text = data.get("text", "")
                theme = data.get("theme")
                language = data.get("language", "en")
                seed = data.get("seed")
        except Exception:
            text = text or ""

        if path == "cv/parse":
            return Response(content=json.dumps(_parse(text)), media_type="application/json")
        if path == "cv/parse_file":
            if uploaded is None:
                return JSONResponse({"ok": False, "error": "no file uploaded (multipart field 'file')"},
                                    status_code=400)
            return Response(content=json.dumps(_parse_file(uploaded)), media_type="application/json")
        if path == "cv/generate":
            res = _generate(text, theme, language, seed)
            return Response(content=json.dumps(res, ensure_ascii=False), media_type="application/json")
        if path == "":
            return _landing()
        return JSONResponse({"ok": False, "error": f"unknown POST route: /{orig}"}, status_code=404)

    return JSONResponse({"ok": False, "error": "method not allowed"}, status_code=405)


app.add_api_route("/{full_path:path}", _dispatch, methods=["GET", "POST"])
